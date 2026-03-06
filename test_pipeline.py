"""End-to-end pipeline test: DOCX + TXT -> AST -> Typst -> PDF"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from api.extractor import extract_text
from api.cleaner import clean_text
from api.section_detector import detect_sections
from api.template_mapper import map_document
from api.typst_generator import generate
from api.typst_compiler import _find_typst
import asyncio, os

print("=" * 60)
print("  PIPELINE END-TO-END TEST")
print("=" * 60)

# ---- Step 0: Verify typst.exe is found ----
try:
    typst_path = _find_typst()
    print(f"[PASS] typst found: {typst_path}")
except FileNotFoundError as e:
    print(f"[FAIL] typst not found: {e}")
    sys.exit(1)

# ---- Step 1: Create a test DOCX with special characters ----
doc = Document()
doc.add_heading("Host Attachment & Fluid Shear are Integrated into a Signal", 0)
doc.add_paragraph("Riddhi Shaha, S.A. Morshed, M.S.I. Khan")
doc.add_heading("Abstract", level=1)
doc.add_paragraph(
    "Attachment to host cells triggers LEE1 promoter induction in a GrlA-dependent manner. "
    "LEE1 is the first transcriptional unit within the LEE region and encodes Ler, "
    "the master regulator of EHEC virulence gene expression. "
    "We used E. coli K12 as a surrogate strain (P < 0.05). "
    "Temperature was 37\u00b0C. Concentration was ~10ng/mL. "
    "Email: test@example.com. Cost was $500. See section #3."
)
doc.add_heading("Introduction", level=1)
doc.add_paragraph(
    "Previous reports show only a moderate induction of LEE1 promoter activity "
    "on exposure to individual environmental cues, but many of these studies were done "
    "using E. coli K12 as a surrogate strain. "
    "The *key finding* is that host_cell attachment is the primary trigger. "
    "See Fig. 1A (19). Reference @ref1 is important."
)
doc.add_heading("Materials and Methods", level=1)
doc.add_paragraph(
    "We infected HeLa epithelial cells with EHEC for 4 h. "
    "Bacteria were grown in LB medium at 37C. "
    "Statistical analysis used Student's t test (P < 0.05)."
)
doc.add_heading("Results", level=1)
doc.add_paragraph(
    "Host-adherent bacteria showed strongly increased LEE1 promoter activity "
    "(~14-fold compared with EHEC grown in LB). "
    "GrlA overexpression led to a hyper-infective phenotype."
)
doc.add_heading("Discussion", level=1)
doc.add_paragraph(
    "Human disease caused by EHEC infection is usually the result of "
    "foodborne transmission. The GrlA-based regulation mechanism is clearly more "
    "complex than a transition from GrlR-bound to unbound states."
)
doc.add_heading("Conclusion", level=1)
doc.add_paragraph(
    "We conclude that host cell attachment triggers LEE1 activation "
    "via a GrlA-dependent mechanism."
)
doc.add_heading("References", level=1)
doc.add_paragraph(
    "[1] Nataro JP, Kaper JB (1998) Diarrheagenic E. coli. Clin Microbiol Rev 11(1): 142-201.\n"
    "[2] McDaniel TK, Kaper JB (1997) A cloned pathogenicity island. MolMicrobiol 23(2):399-407."
)

buf = io.BytesIO()
doc.save(buf)
contents = buf.getvalue()
print(f"[PASS] Test DOCX created ({len(contents)} bytes)")

# ---- Step 2: Extract ----
r1 = extract_text(contents, "paper.docx")
assert r1["status"] == "success", f"Extraction failed: {r1}"
print(f"[PASS] Extracted {r1['word_count']} words")

# ---- Step 3: Clean ----
r2 = clean_text(r1["raw_text"])
assert r2["status"] == "success", f"Cleaning failed: {r2}"
print(f"[PASS] Cleaned: {r2['original_length']} -> {r2['cleaned_length']} chars")

# ---- Step 4: Detect sections ----
r3 = detect_sections(r2["clean_text"])
assert r3["status"] == "success", f"Detection failed: {r3}"
doc_ast = r3["document"]
print(f"[PASS] Sections detected: {r3['section_count']}")
print(f"       Title: {doc_ast['title']}")
print(f"       Authors: {doc_ast['authors']}")
print(f"       Abstract: {doc_ast['abstract'][:60]}...")
for s in doc_ast["sections"]:
    print(f"       - {s['heading']}")

# ---- Step 5: Template mapping ----
for conf in ["ieee", "acm"]:
    for layout in ["single-column", "double-column"]:
        r4 = map_document(doc_ast, conf, layout)
        assert r4["status"] == "success", f"Mapping failed ({conf}/{layout}): {r4}"
        content_blocks = r4["data"]["content"]
        section_types = [b["type"] for b in content_blocks if b["type"] in ("section","subsection","subsubsection")]
        print(f"[PASS] Mapped to {conf}/{layout}: {len(content_blocks)} blocks, {len(section_types)} sections")

# ---- Step 6: Generate Typst + Compile to PDF (IEEE double-column) ----
r5 = generate(document=doc_ast, conference="ieee", layout="double-column")
assert r5["status"] == "success", f"Generation failed: {r5}"
print(f"[PASS] Typst generated: {len(r5['typst_code'])} chars")
print(f"       Output dir: {r5['output_dir']}")

# Check the generated .typ file has proper section/subsection hierarchy
typ_code = r5["typst_code"]
assert 'type: "section"' in typ_code, "Missing section blocks in .typ"
print(f"[PASS] .typ contains section hierarchy")

print(f"[PASS] .typ source generated successfully")

# ---- Step 7: Compile to PDF ----
from api.typst_compiler import compile_typst as compile_fn
r6 = asyncio.run(compile_fn(r5["output_dir"]))
if r6["status"] == "success":
    pdf_size = os.path.getsize(r6["pdf_path"])
    with open(r6["pdf_path"], "rb") as f:
        header = f.read(5)
    assert header == b"%PDF-", f"Not a valid PDF: {header}"
    print(f"[PASS] PDF compiled: {pdf_size:,} bytes ({r6['pdf_path']})")
    print(f"       PDF header: {header}")
else:
    print(f"[FAIL] Compilation failed: {r6.get('error', 'unknown')}")

# Cleanup
generate.__module__  # just to keep reference
from api.typst_generator import cleanup
cleanup(r5["output_dir"])

print()
print("=" * 60)
print("  ALL TESTS PASSED" if r6["status"] == "success" else "  TESTS FAILED")
print("=" * 60)
