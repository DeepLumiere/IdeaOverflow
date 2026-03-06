"""
Input Layer — Text Extraction Module
=====================================
Extracts raw text from uploaded documents.
Supported: PDF, DOCX, TXT, MD
"""

import os
import tempfile
from pathlib import Path


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """
    Extract raw text from file bytes based on the file extension.

    Returns:
        {
            "status": "success",
            "filename": "paper.docx",
            "format": "docx",
            "raw_text": "Deep Learning for OCR\n\nAbstract\nThis paper..."
        }
    """
    ext = Path(filename).suffix.lower()

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".txt": _extract_txt,
        ".md": _extract_markdown,
        ".tex": _extract_txt,  # LaTeX is plain text
    }

    extractor = extractors.get(ext)
    if not extractor:
        return {
            "status": "error",
            "filename": filename,
            "format": ext.lstrip("."),
            "raw_text": "",
            "error": f"Unsupported format: {ext}. Supported: .pdf, .docx, .txt, .md, .tex",
        }

    try:
        raw_text = extractor(file_bytes)
        return {
            "status": "success",
            "filename": filename,
            "format": ext.lstrip("."),
            "raw_text": raw_text,
            "char_count": len(raw_text),
            "word_count": len(raw_text.split()),
        }
    except Exception as e:
        return {
            "status": "error",
            "filename": filename,
            "format": ext.lstrip("."),
            "raw_text": "",
            "error": str(e),
        }


# ── PDF Extraction ──────────────────────────────────────────────────

def _extract_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF using PyMuPDF with font-size-aware,
    column-aware reading order.  Falls back to simple page.get_text()
    if the structured path produces nothing.
    """
    import fitz  # PyMuPDF
    import re as _re

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name

    try:
        doc = fitz.open(tmp_path)
        spans = _extract_spans(doc)
        doc.close()
    finally:
        os.unlink(tmp_path)

    if spans:
        return "\n\n".join(s["text"] for s in spans if s["text"].strip())

    # Fallback: simple plain-text extraction
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp2:
        tmp2.write(file_bytes)
        tmp2_path = tmp2.name
    try:
        doc = fitz.open(tmp2_path)
        parts = []
        for page in doc:
            t = page.get_text("text")
            if t.strip():
                parts.append(t)
        doc.close()
        return "\n\n".join(parts)
    finally:
        os.unlink(tmp2_path)


def _extract_spans(doc) -> list[dict]:
    """
    Walk every page and collect text blocks with font-size metadata,
    bounding-box coordinates, and a column index so two-column papers
    are read in the correct order.
    """
    import re as _re

    spans: list[dict] = []

    for page_num, page in enumerate(doc):
        page_width = page.rect.width
        blocks = page.get_text("dict", flags=0)["blocks"]

        for block in blocks:
            if block.get("type") != 0:          # skip image blocks
                continue

            # Collect all character sizes in this block
            sizes: list[float] = []
            line_texts: list[str] = []
            for line in block.get("lines", []):
                chars_in_line: list[str] = []
                for span in line.get("spans", []):
                    sizes.append(span.get("size", 0.0))
                    chars_in_line.append(span.get("text", ""))
                line_texts.append("".join(chars_in_line))

            text = "\n".join(line_texts).strip()
            if not text:
                continue

            # Clean control characters and ligatures
            text = text.replace("\x00", "")
            text = _re.sub(r"[\x01-\x08\x0b-\x0c\x0e-\x1f\x7f]+", "", text)
            text = text.replace("\ufb01", "fi").replace("\ufb02", "fl").replace("\ufb00", "ff")

            dominant_size = round(max(sizes), 1) if sizes else 0.0
            bbox = block["bbox"]          # (x0, y0, x1, y1)
            x0, y0, x1, y1 = bbox
            width = x1 - x0

            # Column detection: full-width (spanning >65% of page) → col 0
            # otherwise left half → col 1, right half → col 2
            if page_width and width > page_width * 0.65:
                col = 0
            elif page_width and x0 < page_width / 2:
                col = 1
            else:
                col = 2

            spans.append({
                "page": page_num,
                "size": float(dominant_size),
                "text": text,
                "col": col,
                "y0": float(y0),
                "x0": float(x0),
            })

    if not spans:
        return []

    # Sort by reading order: Page → Column → Top-to-Bottom → Left-to-Right
    spans.sort(key=lambda s: (s["page"], s["col"], s["y0"], s["x0"]))

    return spans


# ── DOCX Extraction ────────────────────────────────────────────────

def _extract_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document
    import io

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


# ── TXT / LaTeX Extraction ─────────────────────────────────────────

def _extract_txt(file_bytes: bytes) -> str:
    """Extract text from plain text or LaTeX files."""
    # Try UTF-8 first, then fallback to latin-1
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace")


# ── Markdown Extraction ─────────────────────────────────────────────

def _extract_markdown(file_bytes: bytes) -> str:
    """Extract text from Markdown (returns raw Markdown text)."""
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1", errors="replace")

    return text
